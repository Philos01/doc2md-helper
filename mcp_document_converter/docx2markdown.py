import os
import sys

DEFAULT_TIMEOUT = 300


def docx_to_markdown(docx_path: str, output_dir: str = None):
    """
    将Word文档转换为Markdown格式，支持 .docx 和 .doc 格式
    
    Args:
        docx_path: Word文档路径
        output_dir: 输出目录，如果为None则使用Word文档所在目录
    """
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    
    os.makedirs(output_dir, exist_ok=True)

    docx_timeout = DEFAULT_TIMEOUT

    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    output_path = os.path.join(output_dir, f"{base_name}.md")
    file_ext = os.path.splitext(docx_path)[1].lower()
    
    try:
        # 首先尝试使用 textract 库（支持 .doc 和 .docx）
        try:
            import textract
            print(f"[DEBUG] 使用textract转换Word文档: {docx_path}")
            
            text = textract.process(docx_path).decode('utf-8')
            
            # 简单格式化为Markdown
            markdown_lines = []
            lines = text.split('\n')
            for line in lines:
                stripped = line.strip()
                if stripped:
                    markdown_lines.append(stripped)
                    markdown_lines.append("")
            
            markdown_text = "\n".join(markdown_lines)
            
            # 保存Markdown文件
            with open(output_path, "w", encoding="utf-8") as md_file:
                md_file.write(markdown_text)
            
            print(f"[DEBUG] 转换成功，Markdown文件: {output_path}")
            return output_path
            
        except ImportError:
            print(f"[DEBUG] textract库未安装")
        
        # 如果是 .docx 文件，尝试 mammoth 或 python-docx
        if file_ext == '.docx':
            # 尝试使用mammoth库
            try:
                import mammoth
                print(f"[DEBUG] 使用mammoth转换Word文档: {docx_path}")
                
                with open(docx_path, "rb") as docx_file:
                    result = mammoth.convert_to_markdown(docx_file)
                    markdown_text = result.value
                    
                    # 保存Markdown文件
                    with open(output_path, "w", encoding="utf-8") as md_file:
                        md_file.write(markdown_text)
                    
                    print(f"[DEBUG] 转换成功，Markdown文件: {output_path}")
                    if result.messages:
                        print(f"[DEBUG] 转换警告: {result.messages}")
                    return output_path
                    
            except ImportError:
                print("[DEBUG] mammoth库未安装，尝试使用python-docx")
            
            # 备选方案：使用python-docx
            try:
                from docx import Document
                print(f"[DEBUG] 使用python-docx转换Word文档: {docx_path}")
                
                doc = Document(docx_path)
                markdown_lines = []
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        markdown_lines.append(text)
                        markdown_lines.append("")  # 添加空行分隔段落
                
                # 处理表格
                for table in doc.tables:
                    markdown_lines.append("")
                    for i, row in enumerate(table.rows):
                        cells = [cell.text.strip() for cell in row.cells]
                        markdown_lines.append("| " + " | ".join(cells) + " |")
                        if i == 0:
                            markdown_lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
                    markdown_lines.append("")
                
                markdown_text = "\n".join(markdown_lines)
                
                # 保存Markdown文件
                with open(output_path, "w", encoding="utf-8") as md_file:
                    md_file.write(markdown_text)
                
                print(f"[DEBUG] 转换成功，Markdown文件: {output_path}")
                return output_path
                
            except ImportError:
                print("[DEBUG] python-docx库未安装")
        else:
            print(f"[DEBUG] .doc 文件需要 textract 或 antiword")
        
        # 尝试使用 antiword (Linux/Mac) 或 pywin32 (Windows) 处理 .doc
        if file_ext == '.doc':
            # Windows: 尝试使用 pywin32
            if sys.platform == 'win32':
                try:
                    import win32com.client
                    print(f"[DEBUG] 使用pywin32转换Word文档: {docx_path}")
                    
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    doc = word.Documents.Open(os.path.abspath(docx_path))
                    
                    # 保存为临时 .docx
                    temp_docx = os.path.join(output_dir, f"{base_name}_temp.docx")
                    doc.SaveAs(os.path.abspath(temp_docx), FileFormat=16)
                    doc.Close()
                    word.Quit()
                    
                    # 递归调用处理 .docx
                    return docx_to_markdown(temp_docx, output_dir)
                    
                except ImportError:
                    print("[DEBUG] pywin32未安装")
                except Exception as e:
                    print(f"[DEBUG] pywin32转换失败: {e}")
            
            # Linux/Mac: 尝试使用 antiword
            else:
                try:
                    import subprocess
                    print(f"[DEBUG] 使用antiword转换Word文档: {docx_path}")
                    
                    result = subprocess.run(
                        ["antiword", docx_path],
                        capture_output=True,
                        text=True,
                        timeout=docx_timeout
                    )
                    
                    if result.returncode == 0:
                        # 简单格式化为Markdown
                        markdown_lines = []
                        lines = result.stdout.split('\n')
                        for line in lines:
                            stripped = line.strip()
                            if stripped:
                                markdown_lines.append(stripped)
                                markdown_lines.append("")
                        
                        markdown_text = "\n".join(markdown_lines)
                        
                        with open(output_path, "w", encoding="utf-8") as md_file:
                            md_file.write(markdown_text)
                        
                        print(f"[DEBUG] 转换成功，Markdown文件: {output_path}")
                        return output_path
                    
                except FileNotFoundError:
                    print("[DEBUG] antiword未安装")
        
        # 最后备选：使用pandoc命令行工具
        try:
            import subprocess
            print(f"[DEBUG] 使用pandoc转换Word文档: {docx_path}")
            
            result = subprocess.run(
                ["pandoc", docx_path, "-o", output_path, "--wrap=none"],
                capture_output=True,
                text=True,
                timeout=docx_timeout
            )
            
            if result.returncode == 0:
                print(f"[DEBUG] 转换成功，Markdown文件: {output_path}")
                return output_path
            else:
                print(f"[DEBUG] pandoc转换失败: {result.stderr}")
                
        except (ImportError, FileNotFoundError):
            pass
        
        # 所有方法都失败了，给出明确的提示
        if file_ext == '.doc':
            raise Exception(
                f"旧版 .doc 文件需要安装额外工具才能转换。\n"
                f"推荐方案：\n"
                f"  1. 将文件另存为 .docx 格式再上传\n"
                f"  2. 或安装 textract: pip install textract\n"
                f"  3. Windows用户: pip install pywin32\n"
                f"  4. Linux用户: sudo apt install antiword"
            )
        else:
            raise Exception("请安装mammoth或python-docx或pandoc以支持Word文档转换")
    
    except Exception as e:
        print(f"[ERROR] Word转换失败: {e}")
        import traceback
        traceback.print_exc()
        raise


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("使用方法: python docx2markdown.py <docx文件路径> [输出目录]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    docx_to_markdown(docx_path, output_dir)
